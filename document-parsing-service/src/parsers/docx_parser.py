"""
Parses Microsoft Word documents (.docx) to Markdown.

This parser uses the `python-docx` library to parse the document.
"""

import time
from io import BytesIO
from typing import Dict, Tuple

from docx import Document as DocxDocument


def parse(content: bytes) -> Tuple[str, Dict]:
    start = time.time()
    doc = DocxDocument(BytesIO(content))
    lines: list[str] = []

    for element in doc.element.body:
        if element.tag.endswith("p"):
            para = next(p for p in doc.paragraphs if p._element is element)
            style = para.style.name if para.style else ""
            if style.startswith("Heading"):
                level = 1
                for n in range(1, 7):
                    if f"Heading {n}" in style:
                        level = n
                        break
                lines.append(f"{'#' * level} {para.text}\n")
            else:
                lines.append(f"{para.text}\n")
        elif element.tag.endswith("tbl"):
            table = next(t for t in doc.tables if t._element is element)
            rows_md: list[str] = []
            for r_i, row in enumerate(table.rows):
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows_md.append("| " + " | ".join(cells) + " |")
                if r_i == 0:
                    rows_md.append(
                        "| " + " | ".join(["---"] * len(cells)) + " |"
                    )
            lines.append("\n" + "\n".join(rows_md) + "\n")

    props = doc.core_properties
    metadata = {
        "title": getattr(props, "title", None),
        "author": getattr(props, "author", None),
        "created": (
            getattr(props, "created", None).isoformat()
            if getattr(props, "created", None)
            else None
        ),
        "modified": (
            getattr(props, "modified", None).isoformat()
            if getattr(props, "modified", None)
            else None
        ),
        "processing_ms": int((time.time() - start) * 1000),
    }
    return ("\n".join(lines).strip(), metadata)
