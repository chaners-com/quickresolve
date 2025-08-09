#!/usr/bin/env python3
"""
Document Parsing Service
FastAPI service that parses PDF/DOC/DOCX to Markdown, uploads to S3,
and notifies the embedding service.
"""

import os
import time
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

import boto3
import fitz  # PyMuPDF
import requests
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

app = FastAPI(
    title="Document Parsing Service",
    description="Convert PDF/DOC/DOCX to Markdown and forward to embedding",
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost", "http://localhost:8080"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Environment and clients
S3_ENDPOINT = os.getenv("S3_ENDPOINT")
S3_ACCESS_KEY = os.getenv("S3_ACCESS_KEY")
S3_SECRET_KEY = os.getenv("S3_SECRET_KEY")
S3_BUCKET = os.getenv("S3_BUCKET")

EMBEDDING_SERVICE_URL = os.getenv(
    "EMBEDDING_SERVICE_URL", "http://embedding-service:8001"
)
INGESTION_SERVICE_URL = os.getenv(
    "INGESTION_SERVICE_URL", "http://ingestion-service:8000"
)

s3_client = None


@app.on_event("startup")
def on_startup():
    global s3_client
    s3_client = boto3.client(
        "s3",
        endpoint_url=S3_ENDPOINT,
        aws_access_key_id=S3_ACCESS_KEY,
        aws_secret_access_key=S3_SECRET_KEY,
    )
    # Ensure local parsed directory exists for testing purposes
    Path("parsed").mkdir(parents=True, exist_ok=True)


class ParseRequest(BaseModel):
    s3_key: str
    file_id: int
    workspace_id: int
    original_filename: str


class ParseResponse(BaseModel):
    success: bool
    message: str
    parsed_s3_key: Optional[str] = None
    original_size_bytes: int
    parsed_size_bytes: Optional[int] = None
    processing_time_seconds: float
    extracted_metadata: Optional[dict] = None


@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "service": "document-parsing-service",
        "timestamp": datetime.utcnow().isoformat(),
    }


@app.get("/supported-types")
async def get_supported_types():
    return {
        "supported_formats": [
            {"extension": ".pdf", "description": "PDF documents"},
            {"extension": ".docx", "description": "Microsoft Word documents"},
            {
                "extension": ".doc",
                "description": "Microsoft Word legacy documents",
            },
            {
                "extension": ".md",
                "description": "Markdown files (pass-through)",
            },
        ]
    }


def _set_file_status_failed(file_id: int) -> None:
    try:
        requests.put(
            f"{INGESTION_SERVICE_URL}/files/{file_id}/status",
            params={"status": 3},
            timeout=10,
        )
    except Exception:
        pass


def _extract_text_pdf(content: bytes) -> tuple[str, dict]:
    start = time.time()
    doc = fitz.open(stream=content, filetype="pdf")
    parts: list[str] = []
    for i, page in enumerate(doc, start=1):
        text = page.get_text()
        parts.append(f"\n\n## Page {i}\n\n{text.strip()}\n")
    metadata = {
        "page_count": doc.page_count,
        "title": doc.metadata.get("title") if doc.metadata else None,
        "author": doc.metadata.get("author") if doc.metadata else None,
        "processing_ms": int((time.time() - start) * 1000),
    }
    doc.close()
    return ("".join(parts).strip(), metadata)


def _extract_text_docx(content: bytes) -> tuple[str, dict]:
    start = time.time()
    doc = DocxDocument(BytesIO(content))
    lines: list[str] = []

    # Iterate preserving basic structure (headings, paragraphs, tables)
    for element in doc.element.body:
        if element.tag.endswith("p"):
            para = next(p for p in doc.paragraphs if p._element is element)
            style = para.style.name if para.style else ""
            if style.startswith("Heading"):
                level = 1
                for n in range(1, 7):
                    if f"Heading {n}" in style:
                        level = n
                        break
                lines.append(f"{'#' * level} {para.text}\n")
            else:
                lines.append(f"{para.text}\n")
        elif element.tag.endswith("tbl"):
            table = next(t for t in doc.tables if t._element is element)
            # Convert first row as header; simple markdown table
            rows_md: list[str] = []
            for r_i, row in enumerate(table.rows):
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows_md.append("| " + " | ".join(cells) + " |")
                if r_i == 0:
                    rows_md.append(
                        "| " + " | ".join(["---"] * len(cells)) + " |"
                    )
            lines.append("\n" + "\n".join(rows_md) + "\n")

    props = doc.core_properties
    metadata = {
        "title": getattr(props, "title", None),
        "author": getattr(props, "author", None),
        "created": (
            getattr(props, "created", None).isoformat()
            if getattr(props, "created", None)
            else None
        ),
        "modified": (
            getattr(props, "modified", None).isoformat()
            if getattr(props, "modified", None)
            else None
        ),
        "processing_ms": int((time.time() - start) * 1000),
    }
    return ("\n".join(lines).strip(), metadata)


@app.post("/parse/", response_model=ParseResponse)
async def parse_document(request: ParseRequest):
    started = time.time()

    print(f"Parsing document: {request.s3_key}")

    # 1) Download original content from S3
    try:
        obj = s3_client.get_object(Bucket=S3_BUCKET, Key=request.s3_key)
        content_bytes: bytes = obj["Body"].read()
        original_size = len(content_bytes)
    except Exception as e:
        _set_file_status_failed(request.file_id)
        raise HTTPException(
            status_code=500, detail=f"Failed to download from S3: {e}"
        )

    # 2) Parse to Markdown based on original filename extension
    ext = request.original_filename.lower().strip().split(".")[-1]
    parsed_md = ""
    extracted_metadata: dict = {}
    try:
        if ext == "pdf":
            parsed_md, extracted_metadata = _extract_text_pdf(content_bytes)
        elif ext in ("docx", "doc"):
            parsed_md, extracted_metadata = _extract_text_docx(content_bytes)
        elif ext == "md":
            parsed_md = content_bytes.decode("utf-8", errors="ignore")
            extracted_metadata = {}
        else:
            raise ValueError(f"Unsupported file type: .{ext}")
    except Exception as e:
        _set_file_status_failed(request.file_id)
        raise HTTPException(
            status_code=500, detail=f"Failed to parse document: {e}"
        )

    parsed_bytes = parsed_md.encode("utf-8")

    # 3) Upload parsed markdown to S3 under parsed/{workspace_id}/{file_id}.md
    parsed_s3_key = f"parsed/{request.workspace_id}/{request.file_id}.md"

    try:
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=parsed_s3_key,
            Body=parsed_bytes,
            ContentType="text/markdown",
        )
    except Exception as e:
        _set_file_status_failed(request.file_id)
        raise HTTPException(
            status_code=500, detail=f"Failed to upload parsed file to S3: {e}"
        )

    """
    # 4) Save a local copy for testing: ./parsed/{workspace_id}_{file_id}.md
    try:
        local_path = Path("parsed") /
        f"{request.workspace_id}_{request.file_id}.md"
        local_path.write_bytes(parsed_bytes)
    except Exception:
        # Local save is best-effort;
        # do not fail the request if it writes to S3 successfully
        pass
    """

    # 5) Notify embedding-service to embed the parsed content
    try:
        requests.post(
            f"{EMBEDDING_SERVICE_URL}/embed/",
            json={
                "s3_key": parsed_s3_key,
                "file_id": request.file_id,
                "workspace_id": request.workspace_id,
            },
            timeout=30,
        )
    except Exception as e:
        # Mark as failed since embedding cannot proceed
        _set_file_status_failed(request.file_id)
        raise HTTPException(
            status_code=500, detail=f"Failed to notify embedding-service: {e}"
        )

    duration = time.time() - started

    return ParseResponse(
        success=True,
        message="Parsed and forwarded for embedding",
        parsed_s3_key=parsed_s3_key,
        original_size_bytes=original_size,
        parsed_size_bytes=len(parsed_bytes),
        processing_time_seconds=round(duration, 3),
        extracted_metadata=extracted_metadata or None,
    )


"""
Local testing:
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8005)
"""
