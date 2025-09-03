"""
Parses Microsoft Word documents (.docx) to Markdown.
This parser uses the `python-docx` library to parse the document.
It is fast but does not handle images, tables,
sections or other complex elements.
"""

from io import BytesIO
from typing import List, Tuple


class FastDOCXParser:
    VERSION = "fast-docx-parser-1"

    @staticmethod
    async def parse(content: bytes, context: dict) -> Tuple[str, List[dict]]:
        try:
            from docx import Document as DocxDocument
        except Exception as e:
            raise RuntimeError(
                """python-docx is required for fast DOCX
                parsing but not available."""
            ) from e

        try:
            doc = DocxDocument(BytesIO(content))
            lines: list[str] = []

            for element in doc.element.body:
                if element.tag.endswith("p"):
                    para = next(
                        p for p in doc.paragraphs if p._element is element
                    )
                    style = para.style.name if para.style else ""
                    if style.startswith("Heading"):
                        level = 1
                        for n in range(1, 7):
                            if f"Heading {n}" in style:
                                level = n
                                break
                        lines.append(f"{'#' * level} {para.text}\n")
                    else:
                        lines.append(f"{para.text}\n")
                elif element.tag.endswith("tbl"):
                    table = next(
                        t for t in doc.tables if t._element is element
                    )
                    rows_md: list[str] = []
                    for r_i, row in enumerate(table.rows):
                        cells = [
                            c.text.strip().replace("\n", " ")
                            for c in row.cells
                        ]
                        rows_md.append("| " + " | ".join(cells) + " |")
                        if r_i == 0:
                            rows_md.append(
                                "| " + " | ".join(["---"] * len(cells)) + " |"
                            )
                    lines.append("\n" + "\n".join(rows_md) + "\n")

            markdown = "\n".join(lines).strip()
        except Exception as e:
            raise RuntimeError(f"Fast DOCX parsing failed: {e}")

        images: List[dict] = []
        return markdown, images
